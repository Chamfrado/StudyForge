from io import BytesIO

import pytest
from docx import Document
from fastapi import HTTPException

from app.shared.text_extractor import extract_text_from_file


def test_extracts_utf8_text() -> None:
    text = extract_text_from_file("notes.txt", b"hello study forge")

    assert text == "hello study forge"


def test_rejects_invalid_utf8_text() -> None:
    with pytest.raises(HTTPException) as exc_info:
        extract_text_from_file("notes.txt", b"\xff\xfe\xfa")

    assert exc_info.value.status_code == 400
    assert exc_info.value.detail == "Text file must be encoded in UTF-8."


def test_rejects_unsupported_extension() -> None:
    with pytest.raises(HTTPException) as exc_info:
        extract_text_from_file("notes.exe", b"not allowed")

    assert exc_info.value.status_code == 400
    assert "Supported types" in exc_info.value.detail


def test_extracts_docx_text() -> None:
    document = Document()
    document.add_paragraph("First paragraph")
    document.add_paragraph("Second paragraph")

    buffer = BytesIO()
    document.save(buffer)

    text = extract_text_from_file("notes.docx", buffer.getvalue())

    assert text == "First paragraph\n\nSecond paragraph"


def test_rejects_invalid_pdf_content() -> None:
    with pytest.raises(HTTPException) as exc_info:
        extract_text_from_file("notes.pdf", b"this is not a valid pdf")

    assert exc_info.value.status_code == 400
    assert "Could not extract text from PDF" in exc_info.value.detail
